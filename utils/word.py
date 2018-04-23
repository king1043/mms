#coding:utf-8
import sys
sys.path.append("..")

import os
import win32com
from win32com.client import Dispatch, constants
from docx import Document
from utils.log import log
import numpy as np

class Word():
    def __init__(self, file):
        if file.endswith('doc'):
            self._word = Doc(file)

        elif file.endswith('docx'):
            self._word = Docx(file)

        else:
            raise('%s 格式不支持，仅支持doc 或 docx'%file)

    def get_tables(self):
        return self._word.get_tables()

    def get_table(self, table_pos):
        return self._word.get_table(table_pos)

    def get_data(self, table, rows, clos):
        return self._word.get_data(table, rows, clos)

    def get_datas(self, table):
        return self._word.get_datas(table)

class Doc():
    def __init__(self, file):
        self._word = win32com.client.Dispatch('word.application')
        self._doc = self._word.Documents.Open(file)

    def get_tables(self):
        return self._doc.Tables

    def get_table(self, table_pos):
        return self._doc.Tables[table_pos]

    def get_data(self, table, rows, clos):
        return table.Rows[rows].Cells[clos].Range.Text.replace('', '').strip()

    def get_datas(self, table):
        '''
        @summary:
        ---------
        @param table:
        ---------
        @result: 返回np.array类型  取一列可用data[:,j],取一行可用data[i,:],取某个值可用data[i,j]
        '''

        datas = []
        for row in table.Rows:
            row_datas = []
            for cell in row.Cells:
                data = cell.Range.Text.replace('', '').strip()
                row_datas.append(data)

            datas.append(row_datas)

        return np.array(datas)

class Docx():
    def __init__(self, file):
        self._doc = Document(file)

    def get_tables(self):
        return self._doc.tables

    def get_table(self, table_pos):
        return self._doc.tables[table_pos]

    def get_data(self, table, rows, clos):
        return table.cell(rows, clos).text

    def get_datas(self, table):
        log.info('docx 此方法未实现')
        pass

if __name__ == "__main__":
    word = Word(r'D:\WorkSpace\python\mms\utils\program.doc')
    table = word.get_table(0)
    # data = word.get_data(table, 1, 1)
    data = word.get_datas(table)
    # print(data[1,:])
    print(data[:,0])
